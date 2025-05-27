import unittest
import os
import tempfile
from docx import Document  # 需要安装 python-docx
from src.utils.file_utils import WordDocument


class TestWordDocument(unittest.TestCase):
    """测试WordDocument类的功能"""

    def test_remove_headers_and_footers(self):
        """测试删除页眉和页脚功能"""
        # 创建一个临时目录来存放测试文件
        with tempfile.TemporaryDirectory() as tmpdirname:
            input_path = os.path.join(tmpdirname, 'test_document.docx')

            # 创建一个带有页眉和页脚的测试文档
            doc = Document()

            # 添加首页页眉内容
            section = doc.sections[0]
            header = section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = "这是首页页眉"

            # 添加首页页脚内容
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = "这是首页页脚"

            # 添加新节，并设置不同的页眉和页脚
            new_section = doc.add_section()
            new_section_header = new_section.header
            new_section_footer = new_section.footer

            # 设置新节的页眉和页脚内容
            new_section_header_paragraph = new_section_header.paragraphs[0]
            new_section_header_paragraph.text = "这是新节页眉"

            new_section_footer_paragraph = new_section_footer.paragraphs[0]
            new_section_footer_paragraph.text = "这是新节页脚"

            # 保存文档
            doc.save(input_path)

            # 测试 WordDocument 类的功能
            word_doc = WordDocument(input_path)
            word_doc.open()

            # 删除页眉和页脚
            word_doc.remove_headers()
            word_doc.remove_footers()

            # 保存修改后的文档
            output_path = os.path.join(tmpdirname, 'modified_document.docx')
            word_doc.save(output_path)

            # 验证页眉和页脚是否为空
            modified_doc = Document(output_path)

            for section in modified_doc.sections:
                # 检查页眉是否为空
                self.assertEqual(len(section.header.paragraphs), 0)

                # 检查页脚是否为空
                self.assertEqual(len(section.footer.paragraphs), 0)
