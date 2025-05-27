"""
file_utils.py

提供与文件操作相关的实用函数，包括 Word 文档处理、路径检查等。
"""

from docx import Document


class WordDocument:
    """用于打开和读取Word文档内容的类

    Attributes:
        file_path (str): Word文件的路径
        doc (Document): 文档对象
        content (str): 文档内容
    """

    def __init__(self, file_path):
        """初始化Word文档对象

        Args:
            file_path (str): Word文件的路径
        """
        self.file_path = file_path
        self.doc = None
        self.content = None

    def open(self):
        """打开Word文档并加载内容

        Returns:
            bool: 打开成功返回True，否则返回False
        """
        try:
            self.doc = Document(self.file_path)
            return True
        except FileNotFoundError as e:
            print(f"Error opening file: {e}")
            return False


    def remove_headers(self):
        """清除所有节的主页眉和偶数页页眉内容，并移除多余空段落"""
        try:
            if self.doc:
                for section in self.doc.sections:
                    header = section.header
                    even_header = section.even_page_header  # 获取偶数页页眉
                    header.is_linked_to_previous = False

                    # 清除主页眉内容并删除空段落
                    for paragraph in list(header.paragraphs):
                        # 清除段落中的所有运行（文本、字段等）
                        for run in list(paragraph.runs):
                            run.text = ""
                        # 如果段落为空，则删除该段落
                        if len(paragraph.text.strip()) == 0:
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._element = None

                    # 确保主页眉至少有一个段落用于后续操作
                    if len(header.paragraphs) == 0:
                        header.add_paragraph()

                    # 处理偶数页页眉
                    if even_header:
                        # 清除偶数页页眉内容并删除空段落
                        for paragraph in list(even_header.paragraphs):
                            # 清除段落中的所有运行（文本、字段等）
                            for run in list(paragraph.runs):
                                run.text = ""
                            # 如果段落为空，则删除该段落
                            if len(paragraph.text.strip()) == 0:
                                p = paragraph._element
                                p.getparent().remove(p)
                                p._element = None

                        # 确保偶数页页眉至少有一个段落用于后续操作
                        if len(even_header.paragraphs) == 0:
                            even_header.add_paragraph()

                return True
            return False
        except Exception as e:
            print(f"Error removing headers: {e}")
            return False

    def remove_footers(self):
        """清除所有节的主页脚和偶数页页脚内容，并移除多余空段落"""
        try:
            if self.doc:
                for section in self.doc.sections:
                    footer = section.footer
                    footer.is_linked_to_previous = False

                    # 清除段落内容并删除空段落
                    for paragraph in list(footer.paragraphs):
                        # 删除段落中的所有运行（文本、字段等）
                        for run in list(paragraph.runs):
                            run.text = ""
                        # 删除空段落后保留结构
                        if len(paragraph.text.strip()) == 0:
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._element = None

                    # 确保至少有一个空段落用于后续操作
                    if len(footer.paragraphs) == 0:
                        footer.add_paragraph()

                return True
            return False
        except Exception as e:
            print(f"Error removing footers: {e}")
            return False

    def _add_page_number_to_footer(self, footer):
        """向指定页脚中添加页码字段，使用推荐的标准方式"""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt

        # 清空现有段落（仅清空文本，保留空段落）
        for paragraph in list(footer.paragraphs):
            paragraph.text = ""

        # 添加新段落并居中
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码文本和字段
        run = paragraph.add_run("第")
        run.font.size = Pt(10)

        # 插入 PAGE 字段
        page_field = paragraph.add_run()
        page_field._element.append(OxmlElement("w:fldChar"))
        page_field._element[0].set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'PAGE \\* MERGEFORMAT'
        page_field._element.append(instr_text)

        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        page_field._element.append(end_field)

        run = paragraph.add_run("页/共")
        run.font.size = Pt(10)

        # 插入 NUMPAGES 字段
        pages_field = paragraph.add_run()
        pages_field._element.append(OxmlElement("w:fldChar"))
        pages_field._element[0].set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = 'NUMPAGES \\* MERGEFORMAT'
        pages_field._element.append(instr_text)

        end_field = OxmlElement("w:fldChar")
        end_field.set(qn("w:fldCharType"), "end")
        pages_field._element.append(end_field)

        run = paragraph.add_run("页")
        run.font.size = Pt(10)

    def add_page_numbers(self):
        """在文档页脚添加页码，支持奇偶页不同设置，并确保字段可被识别"""
        try:
            if self.doc:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement

                for section in self.doc.sections:
                    # 获取主（奇数）页脚
                    footer = section.footer
                    footer.is_linked_to_previous = False  # 断开链接

                    # 清空现有段落
                    for paragraph in list(footer.paragraphs):
                        p = paragraph._element
                        p.getparent().remove(p)
                        p._element = None

                    # 添加页码
                    self._add_page_number_to_footer(footer)

                    # 如果存在偶数页页脚，也添加页码
                    even_footer = section.even_page_footer
                    if even_footer:
                        # 断开链接
                        even_footer.is_linked_to_previous = False

                        # 清空现有段落
                        for paragraph in list(even_footer.paragraphs):
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._element = None

                        self._add_page_number_to_footer(even_footer)

                # 插入一个空白节以帮助 Word 刷新字段
                self.doc.add_section()

                return True
        except Exception as e:
            print(f"Error adding page numbers: {e}")
            return False

    def save(self, output_path=None):
        """保存文档到指定路径

        Args:
            output_path (str, optional): 输出文件路径。如果未指定，将覆盖原始文件。
                Defaults to None.

        Returns:
            bool: 保存成功返回True，否则返回False
        """
        try:
            if self.doc:
                # 如果未指定输出路径，则使用原始文件路径
                if output_path is None:
                    output_path = self.file_path

                # 保存文档
                self.doc.save(output_path)
                return True
            return False
        except Exception as e:
            print(f"Error saving file: {e}")
            return False
